"""Парсинг hard + soft job skills из DOCX регламентов."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parents[1]

from scripts.sync_global_regulations_from_sources import (  # noqa: E402
    SOFT_SKILL_RANK_BASE,
    _find_skill_tables,
    _find_skills_table,
    collect_parsed,
    load_url_maps,
)


def _docx_for_position(position_code: str) -> Path | None:
    by_pos = collect_parsed(load_url_maps())
    row = by_pos.get(position_code)
    if not row:
        return None
    return Path(row.source_file)


@pytest.mark.skipif(
    _docx_for_position("DOC_INPATIENT") is None,
    reason="docs/regulations not present locally",
)
def test_mmc_docx_splits_hard_and_soft_tables():
    path = _docx_for_position("DOC_INPATIENT")
    assert path is not None
    hard, soft = _find_skill_tables(Document(str(path)))
    assert len(hard) == 7
    assert len(soft) == 7
    assert "обход" in hard[0][1].lower()
    assert "коммуника" in soft[0][1].lower()


@pytest.mark.skipif(
    _docx_for_position("DOC_INPATIENT") is None,
    reason="docs/regulations not present locally",
)
def test_mmc_combined_ranks_hard_then_soft():
    path = _docx_for_position("DOC_INPATIENT")
    assert path is not None
    skills = _find_skills_table(Document(str(path)))
    assert len(skills) == 14
    assert [r for r, _ in skills[:7]] == list(range(1, 8))
    assert [r for r, _ in skills[7:]] == list(range(SOFT_SKILL_RANK_BASE, SOFT_SKILL_RANK_BASE + 7))
    assert skills[7][1] == _find_skill_tables(Document(str(path)))[1][0][1]


@pytest.mark.skipif(
    _docx_for_position("HR_MANAGER") is None,
    reason="docs/regulations not present locally",
)
def test_default_docx_single_skills_table_unchanged():
    path = _docx_for_position("HR_MANAGER")
    assert path is not None
    hard, soft = _find_skill_tables(Document(str(path)))
    assert len(hard) == 7
    assert soft == []
    skills = _find_skills_table(Document(str(path)))
    assert len(skills) == 7
    assert [r for r, _ in skills] == list(range(1, 8))
