"""
Собрать app/data/universal/regulations_enrichment.json из DOCX в docs/regulations/.
Запуск из корня репозитория (нужен пакет python-docx):
  pip install python-docx
  python scripts/build_regulations_enrichment_json.py
"""

from __future__ import annotations

import json
import re
from pathlib import Path

try:
    from docx import Document
except ImportError:
    raise SystemExit("Установите: pip install python-docx")

ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "docs" / "regulations"
OUT = ROOT / "app" / "data" / "universal" / "regulations_enrichment.json"

SECTION_RE = re.compile(r"^(\d+)\.\s+")


def _meta_table(doc) -> dict[str, str]:
    meta: dict[str, str] = {}
    if not doc.tables:
        return meta
    for row in doc.tables[0].rows:
        cells = [c.text.replace("\n", " ").strip() for c in row.cells]
        if len(cells) >= 2 and cells[0]:
            meta[cells[0]] = cells[1]
    return meta


def _paragraphs(doc) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def _sections(paras: list[str]) -> dict[int, list[str]]:
    """Номер раздела (1,2,3,…) → абзацы до следующего раздела с номером."""
    sections: dict[int, list[str]] = {}
    current: int | None = None
    for line in paras:
        m = SECTION_RE.match(line)
        if m:
            current = int(m.group(1))
            sections.setdefault(current, []).append(line)
            continue
        if current is not None:
            sections[current].append(line)
    return sections


def _find_url(texts: list[str]) -> str | None:
    for t in texts:
        for m in re.finditer(r"https?://[^\s\)\]\"']+", t):
            return m.group(0)[:512]
    return None


def extract_one(path: Path) -> dict | None:
    doc = Document(str(path))
    meta = _meta_table(doc)
    code = (meta.get("Код регламента") or "").strip()
    if not code:
        return None
    paras = _paragraphs(doc)
    sec = _sections(paras)
    goal_parts = sec.get(2, [])
    goal_summary = " ".join(goal_parts[1:])[:512] if len(goal_parts) > 1 else None
    s3 = sec.get(3, [])
    ckp_short = (s3[1] if len(s3) > 1 else (s3[0] if s3 else None)) or None
    if ckp_short:
        ckp_short = ckp_short[:512]
    ckp_chunks = []
    for n in (3, 4, 5):
        if n in sec:
            ckp_chunks.extend(sec[n])
    ckp_full = "\n\n".join(ckp_chunks) if ckp_chunks else None
    url = _find_url(paras)
    name = (meta.get("Должность") or "").replace("\n", " ")
    name = re.sub(r"\s*\([A-Z][A-Z0-9_]*\)\s*", "", name).strip()
    regulation_name = (f"Регламент: {name}" if name else None)
    if regulation_name:
        regulation_name = regulation_name.replace("|", " ").strip()[:256]
    return {
        "regulation_code": code,
        "regulation_name": regulation_name,
        "goal_summary": goal_summary,
        "ckp_short": ckp_short,
        "ckp_full": ckp_full,
        "google_doc_url": url,
    }


def main() -> None:
    if not DOC_DIR.is_dir():
        print("Нет папки", DOC_DIR)
        return
    regulations: list[dict] = []
    for p in sorted(DOC_DIR.glob("*.docx")):
        row = extract_one(p)
        if row:
            regulations.append(row)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    OUT.write_text(json.dumps({"regulations": regulations}, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {len(regulations)} rows to {OUT}")


if __name__ == "__main__":
    main()
