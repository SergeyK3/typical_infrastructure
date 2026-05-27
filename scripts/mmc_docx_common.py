"""Общие утилиты DOCX для медицинских регламентов ММЦ."""

from __future__ import annotations

import copy
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

SKILLS_VERSION = "V1 от 26.05.2026"

SKILLS_SECTION_INTRO = (
    "Информация по навыкам может устаревать. Актуальную версию смотрите "
    "в своих должностных регламентах и матрицах компетенций."
)
SKILLS_SECTION_INTRO2 = "После обновления навыков следует детализировать целевой ориентир."

# Hard skills (профстандарты, клиника) — прежний SKILLS_BY_CODE
HARD_SKILLS_BY_CODE: dict[str, list[str]] = {
    "ADM_ZAM_LECH": [
        "Клиническое руководство амбулаторным и стационарным контуром",
        "Внедрение и контроль клинических протоколов онкопомощи",
        "Организация медико-экспертной оценки качества и разборов инцидентов",
        "Управление врачебным и сестринским персоналом (графики, компетенции)",
        "Планирование ресурсов и загрузки клинических подразделений",
        "Взаимодействие с контролирующими органами и профессиональными сообществами",
        "Стратегическое развитие медицинского блока и непрерывность маршрута пациента",
    ],
    "ADM_ZAMADM": [
        "Организация АХО и эксплуатации здания медицинского центра",
        "Управление закупками, договорами и подрядчиками",
        "Административный документооборот и контроль исполнения поручений",
        "Организация пропускного режима, охраны и логистики",
        "Бюджетирование и контроль расходов административного блока",
        "Координация внутренних сервисов для клинических подразделений",
        "Антикризисное управление форс-мажорами без остановки медпомощи",
    ],
    "ADM_ZAM_AMBUL": [
        "Организация потоков пациентов и графиков амбулаторного приёма",
        "Контроль качества амбулаторной онкопомощи и онкоскрининга",
        "Управление врачами и медсёстрами амбулатории",
        "Маршрутизация в стационар, диагностику и смежные службы",
        "Анализ KPI амбулаторного контура и инициирование улучшений",
        "Диспансерное наблюдение и преемственность со стационаром",
        "Работа с жалобами и обратной связью пациентов амбулатории",
    ],
    "ADM_ZAM_QUAL": [
        "Планирование и проведение внутренних клинических аудитов",
        "Методология root cause analysis и CAPA",
        "Мониторинг соблюдения протоколов и инфекционного контроля",
        "Подготовка отчётности по качеству для руководства и проверок",
        "Организация обучения персонала по качеству и безопасности",
        "Экспертиза медицинской документации и клинических решений",
        "Внедрение показателей качества и best practices",
    ],
    "DOC_AMBUL": [
        "Сбор анамнеза и клинический осмотр онкопациента",
        "Назначение и интерпретация диагностических обследований",
        "Постановка диагноза и выбор тактики по клиническим рекомендациям",
        "Онкоскрининг, раннее выявление и диспансерное наблюдение",
        "Ведение медицинской документации и медицинской информационной системы",
        "Информированное согласие и коммуникация с пациентом и родственниками",
        "Маршрутизация в стационар, консилиумы и смежные специалисты",
    ],
    "DOC_INPATIENT": [
        "Ежедневная оценка состояния и обходы стационарных пациентов",
        "Назначение и коррекция стационарной противоопухолевой терапии",
        "Неотложная помощь и эскалация ухудшений состояния",
        "Координация медсестёр и младшего персонала на посту",
        "Ведение стационарной медицинской документации",
        "Соблюдение инфекционного контроля и безопасности пациента",
        "Оформление выписки и преемственность с амбулаторным звеном",
    ],
    "HEAD_DEPT": [
        "Организация работы госпитального отделения и графиков дежурств",
        "Контроль качества стационарной помощи и клинических протоколов",
        "Управление врачебным и сестринским персоналом отделения",
        "Разбор сложных случаев, осложнений и инцидентов",
        "Планирование загрузки коек и ресурсов отделения",
        "Взаимодействие с диагностикой, ОР, реанимацией, амбулаторией",
        "Анализ KPI отделения и инициирование улучшений процессов",
    ],
    "NURSE_AMBUL": [
        "Подготовка кабинета и инструментария к амбулаторному приёму",
        "Измерение vital signs, ЭКГ и манипуляции по назначению врача",
        "Венепункция, инъекции и забор биоматериала",
        "Маркировка образцов и передача в лабораторию",
        "Ведение процедурной документации и учёта медизделий",
        "Инфекционный контроль и асептика в кабинете",
        "Коммуникация с пациентом и инструктаж по подготовке к процедурам",
    ],
    "WARD_NURSE": [
        "Выполнение лекарственных назначений и инфузионной терапии",
        "Мониторинг vital signs, боли и рисков падений",
        "Сестринский уход, гигиена и помощь пациентам с ограниченной мобильностью",
        "Профилактика пролежней и позиционирование пациентов",
        "Своевременная эскалация изменений состояния врачу",
        "Ведение листов назначений и сменной отчётности",
        "Инфекционный контроль и правила хранения лекарственных средств",
    ],
    "DEPT_CHIEF_NURSE": [
        "Планирование графиков и укомплектованность смен",
        "Контроль стандартов сестринского ухода на постах",
        "Наставничество и инструктаж медсестёр и младшего персонала",
        "Организация инфекционного контроля в отделении",
        "Учёт медизделий, расходников и участие в инвентаризациях",
        "Разбор инцидентов с участием медсестёр и внедрение CAPA",
        "Оперативная связь с заведующим отделением и врачами",
    ],
    "CALL_REG": [
        "Приём и маршрутизация входящих обращений пациентов",
        "Запись на приём с учётом правил центра и загрузки расписания",
        "Консультирование по услугам, подготовке и документам",
        "Ведение карточек обращений в CRM/МИС без ошибок",
        "Соблюдение конфиденциальности персональных данных",
        "Работа по скриптам и нормативам SLA колл-центра",
        "Эскалация сложных обращений руководителю или клинике",
    ],
    "CALL_OUTBOUND": [
        "Исходящие кампании по скриптам (напоминания, скрининг, reactivation)",
        "Работа с возражениями и фиксация причин отказов",
        "Конверсия в запись на приём и передача в регистратуру",
        "Дисциплина обзвона и соблюдение временных окон",
        "Ведение статусов в CRM и качество данных",
        "Этика медицинского обзвона и соблюдение законодательства",
        "Эскалация сложных кейсов старшему оператору",
    ],
    "NURSE_PROCEDURE": [
        "Венепункция, инъекции и инфузионная терапия",
        "Подготовка пациента к химиотерапии и процедурам по чек-листам",
        "Работа с цитостатиками и соблюдение правил безопасности",
        "Забор крови, маркировка и транспортировка биоматериала",
        "Асептика, инфекционный контроль и утилизация отходов Б/В",
        "Ведение журналов процедур и учёта медизделий",
        "Остановка процедуры при риске для пациента",
    ],
    "NURSE_HOUSEKEEP": [
        "Организация стирки, выдачи и учёта белья",
        "Контроль санобработки палат и постов по графику",
        "Постановка задач санитаркам и проверка качества уборки",
        "Учёт моющих и дезсредств, заявки в АХО",
        "Соблюдение режимов дезинфекции и инфекционного контроля",
        "Обход палат и оперативное устранение замечаний",
        "Взаимодействие со старшей медсестрой и палатной службой",
    ],
    "ORDERLY": [
        "Влажная уборка палат, коридоров и санузлов по чек-листам",
        "Смена белья и соблюдение графика уборки",
        "Классификация и вынос медицинских отходов",
        "Применение дезинфектантов по инструкциям и концентрациям",
        "Соблюдение СИЗ и гигиены рук",
        "Немедленное сообщение о рисках инфекции и поломках",
        "Уважительное обращение с пациентами в зоне уборки",
    ],
}

# Soft job skills — прикладные социально-психологические навыки (docs/skill_assessment/soft_job_skills.md)
SOFT_SKILLS_BY_CODE: dict[str, list[str]] = {
    "ADM_ZAM_LECH": [
        "Executive presence и удержание клинического авторитета",
        "Фасилитация разборов и консенсуса в медицинской команде",
        "Дипломатия при междисциплинарных и организационных конфликтах",
        "Управление доверием врачей и медсестёр",
        "Контекстуальное понимание организационной динамики центра",
        "Стрессоустойчивость при инцидентах и проверках",
        "Адаптивная коммуникация с руководством и контролирующими органами",
    ],
    "ADM_ZAMADM": [
        "Переговоры с подрядчиками и внутренними заказчиками",
        "Координация ожиданий между АХО и клиническими подразделениями",
        "Деэскалация конфликтов в условиях дефицита ресурсов",
        "Политическая чувствительность в административной среде",
        "Управление впечатлением при форс-мажорах",
        "Стрессоустойчивость и самоконтроль",
        "Согласование приоритетов при неопределённости сроков",
    ],
    "ADM_ZAM_AMBUL": [
        "Адаптивная коммуникация с пациентами и очередями",
        "Убеждение и мотивация команды амбулатории без давления",
        "Согласование ожиданий между врачами, медсестрами и регистратурой",
        "Социальное восприятие узких мест маршрута пациента",
        "Удержание доверия при жалобах и задержках",
        "Стрессоустойчивость при пиковой нагрузке",
        "Работа с неопределённостью при перегрузке расписания",
    ],
    "ADM_ZAM_QUAL": [
        "Беспристрастная социальная диагностика в аудитах",
        "Конструктивное влияние без прямого давления (деэскалация сопротивления)",
        "Фасилитация разборов инцидентов и CAPA",
        "Управление доверием к экспертизе качества",
        "Перенос критики и работа с конфликтом интересов",
        "Адаптивная коммуникация при обучении персонала",
        "Толерантность к неоднозначности клинических кейсов",
    ],
    "DOC_AMBUL": [
        "Активное слушание и распознавание эмоций пациента",
        "Адаптивная коммуникация при сообщении диагноза и тактики лечения",
        "Деэскалация конфликтов и удержание доверия (trust management)",
        "Эмпатия и прогнозирование реакции пациента и родственников",
        "Стрессоустойчивость и самоконтроль при высокой нагрузке",
        "Координация ожиданий со смежными специалистами",
        "Работа с неопределённостью до получения результатов обследований",
    ],
    "DOC_INPATIENT": [
        "Спокойная коммуникация при ухудшении состояния и с тревожными родственниками",
        "Своевременная эскалация с сохранением доверия команды",
        "Координация медсестёр и младшего персонала без микроменеджмента",
        "Распознавание напряжения в палате и групповой динамики",
        "Стрессоустойчивость в ночные смены и неотложных ситуациях",
        "Согласование тактики с консилиумом и заведующим",
        "Перенос критики и разборов без оборонительной реакции",
    ],
    "HEAD_DEPT": [
        "Лидерство и удержание авторитета в критических ситуациях",
        "Фасилитация разборов осложнений и инцидентов",
        "Дипломатия между врачами, медсестрами и администрацией",
        "Мотивационный анализ поведения сотрудников",
        "Стрессоустойчивость при перегрузке отделения",
        "Согласование ожиданий с амбулаторным и диагностическим контуром",
        "Организационная навигация внутри медицинского блока",
    ],
    "NURSE_AMBUL": [
        "Доброжелательная адаптивная коммуникация с тревожными пациентами",
        "Активное слушание при инструктаже и подготовке к процедурам",
        "Считывание дискомфорта и невербальных сигналов",
        "Согласование действий с врачом при отклонениях",
        "Стрессоустойчивость в высоком темпе приёма",
        "Удержание доверия при болезненных манипуляциях",
        "Спокойная деэскалация при отказе от процедур",
    ],
    "WARD_NURSE": [
        "Эмпатия и поддержка пациентов с ограниченной мобильностью",
        "Спокойная коммуникация при боли и страхе",
        "Своевременная эскалация врачу без паники",
        "Координация с коллегами на посту в смене",
        "Стрессоустойчивость в ночные дежурства",
        "Распознавание риска падений и агрессии",
        "Адаптация стиля общения к состоянию пациента",
    ],
    "DEPT_CHIEF_NURSE": [
        "Наставничество и конструктивная обратная связь",
        "Деэскалация конфликтов в смене",
        "Координация медсестёр и санитарного звена",
        "Удержание дисциплины стандартов без авторитарности",
        "Стрессоустойчивость при дефиците кадров",
        "Согласование ожиданий с заведующим и врачами",
        "Социальное восприятие напряжения на постах",
    ],
    "CALL_REG": [
        "Активное слушание и conversational intelligence",
        "Адаптивная коммуникация с тревожными и конфликтными пациентами",
        "Деэскалация претензий по записи и ожиданию",
        "Удержание доверия к центру в телефонном контакте",
        "Стрессоустойчивость при потоке звонков",
        "Точная фиксация ожиданий и следующего шага",
        "Согласование эскалации без потери лица для пациента",
    ],
    "CALL_OUTBOUND": [
        "Работа с возражениями без агрессии (adaptive communication)",
        "Убеждение и удержание доверия в холодном контакте",
        "Распознавание эмоций и скрытых мотивов отказа",
        "Стрессоустойчивость к массовым отказам",
        "Дисциплина фиксации с эмпатичным тоном",
        "Этичное влияние в рамках скрипта",
        "Согласование передачи в регистратуру без разрыва контакта",
    ],
    "NURSE_PROCEDURE": [
        "Спокойная поддержка пациентов перед химиотерапией",
        "Объяснение процедур понятным языком",
        "Деэскалация страха и отказов от манипуляций",
        "Согласование с врачом при риске для пациента",
        "Стрессоустойчивость при работе с цитостатиками",
        "Внимание к невербальным сигналам дискомфорта",
        "Удержание доверия при повторных венепункциях",
    ],
    "NURSE_HOUSEKEEP": [
        "Уважительное взаимодействие с пациентами во время уборки",
        "Спокойная постановка задач санитаркам",
        "Деэскалация претензий к санитарному состоянию",
        "Согласование приоритетов с палатной службой",
        "Настойчивость без конфликта при нарушениях санитарии",
        "Стрессоустойчивость при множестве замечаний",
        "Чёткая эскалация рисков инфекции",
    ],
    "ORDERLY": [
        "Тактичное общение с пациентами в палате",
        "Соблюдение дистанции и уважение к уязвимости больных",
        "Спокойная реакция на замечания",
        "Своевременное сообщение о рисках без паники",
        "Исполнительность и согласование с сестрой-хозяйкой",
        "Стрессоустойчивость при физической нагрузке",
        "Аккуратность в конфликтных ситуациях в коридорах",
    ],
}

# Обратная совместимость для patch_mmc_regulations_skills_section.py
SKILLS_BY_CODE = HARD_SKILLS_BY_CODE


def block_element(block) -> CT_P | CT_Tbl:
    if isinstance(block, Paragraph):
        return block._p
    if isinstance(block, Table):
        return block._tbl
    raise TypeError(type(block))


def compact_paragraph(paragraph: Paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def compact_table_paragraphs(table: Table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                compact_paragraph(paragraph)


def set_table_borders_all(table: Table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    for child in list(tbl_pr):
        if child.tag == qn("w:tblBorders"):
            tbl_pr.remove(child)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        side = OxmlElement(f"w:{edge}")
        side.set(qn("w:val"), "single")
        side.set(qn("w:sz"), "4")
        side.set(qn("w:space"), "0")
        side.set(qn("w:color"), "auto")
        borders.append(side)
    tbl_pr.append(borders)


def format_regulation_tables(doc: Document, tables: Iterable[Table] | None = None) -> None:
    targets = list(tables) if tables is not None else list(doc.tables)
    for table in targets:
        set_table_borders_all(table)
        compact_table_paragraphs(table)


def set_cell_text(cell: _Cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        compact_paragraph(paragraph)


def fill_skills_table(table: Table, skills: list[str]) -> None:
    rows = [["№", "Навык", "Приоритет", "Версия"]]
    for i, title in enumerate(skills, start=1):
        rows.append([str(i), title, str(i), SKILLS_VERSION])
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    for ri, row_data in enumerate(rows):
        if ri >= len(table.rows):
            table.add_row()
        row = table.rows[ri]
        for ci, value in enumerate(row_data):
            if ci < len(row.cells):
                set_cell_text(row.cells[ci], value)


def insert_paragraph_before(block, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    block_element(block).addprevious(new_p)
    paragraph = Paragraph(new_p, block._parent)
    if text:
        paragraph.add_run(text)
    compact_paragraph(paragraph)
    return paragraph


def insert_paragraph_after(block, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    block_element(block).addnext(new_p)
    paragraph = Paragraph(new_p, block._parent)
    if text:
        paragraph.add_run(text)
    compact_paragraph(paragraph)
    return paragraph


def renumber_sections_from(doc: Document, start: int = 7) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        match = re.match(r"^(\d+)\.\s+(.*)$", text)
        if not match:
            continue
        number = int(match.group(1))
        if number >= start:
            paragraph.text = f"{number + 1}. {match.group(2)}"


def find_kpi_heading(doc: Document) -> Paragraph:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if re.match(r"^\d+\.\s+", text) and "KPI" in text.upper():
            return paragraph
    raise ValueError("KPI heading not found")


def insert_table_before(doc: Document, block, rows: int, cols: int) -> Table:
    table = doc.add_table(rows=rows, cols=cols)
    block_element(block).addprevious(table._tbl)
    return table


def has_mmc_ab_skills_section(doc: Document) -> bool:
    has_heading = has_a = has_b = False
    for para in doc.paragraphs:
        t = para.text.strip()
        low = t.lower()
        if t.startswith("7.") and "навыки" in low:
            has_heading = True
        if (low.startswith("a.") or low.startswith("а.")) and "hard" in low:
            has_a = True
        if (low.startswith("b.") or low.startswith("б.")) and "soft" in low:
            has_b = True
    return has_heading and has_a and has_b


def position_code_from_filename(name: str) -> str | None:
    stem = Path(name).stem
    for code in sorted(HARD_SKILLS_BY_CODE, key=len, reverse=True):
        if stem.endswith(f"_{code}"):
            return code
    return None


def insert_mmc_skills_before_kpi(doc: Document, position_code: str) -> bool:
    """Вставить раздел 7 (A hard + B soft) перед KPI; сдвинуть нумерацию 7→8 …"""
    if position_code not in HARD_SKILLS_BY_CODE:
        return False
    if has_mmc_ab_skills_section(doc):
        return False

    hard = HARD_SKILLS_BY_CODE[position_code]
    soft = SOFT_SKILLS_BY_CODE[position_code]
    kpi_heading = find_kpi_heading(doc)
    renumber_sections_from(doc, 7)

    insert_paragraph_before(kpi_heading, "7. Ключевые навыки")
    insert_paragraph_before(kpi_heading, SKILLS_SECTION_INTRO)
    insert_paragraph_before(kpi_heading, SKILLS_SECTION_INTRO2)
    insert_paragraph_before(kpi_heading, "A. hard skills")
    hard_table = insert_table_before(doc, kpi_heading, rows=1, cols=4)
    fill_skills_table(hard_table, hard)
    format_regulation_tables(doc, [hard_table])

    para_b = insert_paragraph_after(hard_table, "Б. soft job skills")
    soft_xml = copy.deepcopy(hard_table._tbl)
    para_b._p.addnext(soft_xml)
    soft_table = None
    for tbl in doc.tables:
        if tbl._tbl is soft_xml:
            soft_table = tbl
            break
    if soft_table is None:
        raise ValueError("Soft skills table not found after insert")
    fill_skills_table(soft_table, soft)
    format_regulation_tables(doc, [soft_table])
    return True


def find_skills_table(doc: Document) -> Table:
    for table in doc.tables:
        if not table.rows:
            continue
        hdr = [c.text.strip().lower() for c in table.rows[0].cells]
        if "навык" in " ".join(hdr) and "приоритет" in " ".join(hdr):
            return table
    raise ValueError("Skills table not found")


def setup_section7_skills(doc: Document, position_code: str) -> Table:
    """Подразделы А/Б, две таблицы навыков; возвращает таблицу soft skills."""
    hard = HARD_SKILLS_BY_CODE[position_code]
    soft = SOFT_SKILLS_BY_CODE[position_code]
    hard_table = find_skills_table(doc)

    intro1_para = None
    intro2_para = None
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        if "после обновления навыков" in t.lower():
            intro2_para = para
        elif intro1_para is None and "информация по навыкам" in t.lower():
            intro1_para = para

    if intro1_para is not None:
        intro1_para.text = SKILLS_SECTION_INTRO
        compact_paragraph(intro1_para)
    if intro2_para is not None:
        intro2_para._p.getparent().remove(intro2_para._p)

    insert_paragraph_before(hard_table, SKILLS_SECTION_INTRO2)
    insert_paragraph_before(hard_table, "А. hard skills")
    fill_skills_table(hard_table, hard)

    para_b = insert_paragraph_after(hard_table, "Б. soft job skills")
    soft_xml = copy.deepcopy(hard_table._tbl)
    para_b._p.addnext(soft_xml)
    soft_table = None
    for tbl in doc.tables:
        if tbl._tbl is soft_xml:
            soft_table = tbl
            break
    if soft_table is None:
        raise ValueError("Soft skills table not found after insert")
    fill_skills_table(soft_table, soft)
    format_regulation_tables(doc, [soft_table])
    return soft_table


def set_paragraph_after_heading(doc: Document, heading_prefix: str, needle: str, text: str) -> None:
    for para in doc.paragraphs:
        t = para.text.strip()
        if t.startswith(heading_prefix) and needle.upper() in t.upper():
            insert_paragraph_after(para, text)
            return
    raise ValueError(f"Heading not found: {heading_prefix} {needle}")

