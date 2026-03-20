"""Собрать строки шаблонов KPI из DOCX-регламентов в docs/regulations → kpi_templates1.xlsx."""

from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document

REG_DIR = Path(__file__).resolve().parents[1] / "docs" / "regulations"
OUT_PATH = REG_DIR / "kpi_templates1.xlsx"
OUT_DOWNLOADS = Path.home() / "Downloads" / "kpi_templates1.xlsx"
SHEET = "kpi_templates"

COLS = [
    "kpi_code",
    "kpi_name",
    "unit",
    "period_type",
    "formula_or_rule",
    "default_target",
    "is_active",
    "position_code",
    "primary_dept_type_code",
    "position_name_ru",
]


def _table_meta(table) -> dict[str, str]:
    meta: dict[str, str] = {}
    for row in table.rows:
        cells = [c.text.replace("\n", " ").strip() for c in row.cells]
        if len(cells) >= 2 and cells[0]:
            meta[cells[0]] = cells[1]
    return meta


def _position_code(meta: dict[str, str], regulation_code: str) -> str:
    duty = meta.get("Должность", "")
    m = re.search(r"\(([A-Z][A-Z0-9_]*)\)", duty)
    if m:
        return m.group(1)
    m2 = re.match(r"REG_(.+)_V\d+$", regulation_code)
    return m2.group(1) if m2 else regulation_code


def _position_name_ru(meta: dict[str, str]) -> str:
    duty = meta.get("Должность", "").replace("\n", " ")
    duty = re.sub(r"\s*\([A-Z][A-Z0-9_]*\)\s*", "", duty)
    duty = re.sub(r"\s*\|\s*", " ", duty)
    return duty.strip() or meta.get("Должность", "").strip()


def _dept_type(subdivision: str, position_code: str) -> str:
    s = (subdivision or "").upper()
    if "HR" in s or "КАДР" in subdivision:
        return "HR"
    if "БУХГАЛТ" in subdivision or "ФИНАНС" in subdivision:
        return "ACC"
    if "IT" in s or "ИНФОРМАЦИОНН" in subdivision:
        return "IT"
    if "MKT" in s or "МАРКЕТ" in subdivision:
        return "MKT"
    if "LEAD" in s or "ЛИДОГЕН" in subdivision:
        return "LEAD"
    if "ПРОДАЖ" in subdivision or "КОММЕРЧЕСК" in subdivision:
        return "SALES"
    if "ПРОИЗВОДСТВ" in subdivision or "ТЕХНИЧЕСК" in subdivision:
        return "PROD"
    if "КАЧЕСТВ" in subdivision or "ОКК" in subdivision:
        return "QUAL"
    if "PR " in s or "КОММУНИКАЦ" in subdivision or "СМИ" in subdivision:
        return "PR"
    if "АДМИНИСТР" in subdivision or "УПРАВЛЕНИЕ" in subdivision:
        return "ADM"
    if "ОБЩЕЕ РУКОВОДСТВО" in subdivision or "УПРАВЛЯЮЩ" in subdivision:
        return "ADM"
    # fallback по коду должности
    if position_code.startswith("HR_"):
        return "HR"
    if position_code.startswith("ACC_"):
        return "ACC"
    if position_code.startswith("SALES_"):
        return "SALES"
    if position_code.startswith("QUAL_"):
        return "QUAL"
    if position_code.startswith("MKT_"):
        return "MKT"
    if position_code.startswith("LEAD"):
        return "LEAD"
    if position_code.startswith("PR_"):
        return "PR"
    if "PROD" in position_code or "TECH" in position_code:
        return "PROD"
    if position_code in ("SYSADMIN", "INFO_SYSTEM_SUPPORT"):
        return "IT"
    if position_code == "DIRECTOR":
        return "ADM"
    return "ADM"


def _infer_unit(target: str, how: str) -> str:
    t, h = (target or "").lower(), (how or "").lower()
    if "дн" in h or "time-to-fill" in h or "срок" in h and "реакц" in h:
        return "дней"
    if "шт" in h or "количеств" in h:
        return "шт"
    if "%" in t or "дол" in h or "план" in h and "%" in t:
        return "%"
    if "выручк" in h or "прибыл" in h:
        return "%"
    return "индекс"


def _infer_default_target(target: str) -> float | None:
    if not target:
        return None
    if re.search(r"0\s+критич", target, re.I):
        return 0.0
    nums = re.findall(r"(\d+(?:[.,]\d+)?)\s*%", target)
    if nums:
        v = nums[-1].replace(",", ".")
        try:
            return float(v)
        except ValueError:
            pass
    m = re.search(r"(\d+)\s*[-–]\s*(\d+)\s*%", target)
    if m:
        return (float(m.group(1)) + float(m.group(2))) / 2.0
    if re.search(r"100\s*%", target):
        return 100.0
    return None


def _find_kpi_table(doc: Document):
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        hdr = [c.text.strip().lower() for c in tbl.rows[0].cells]
        joined = " ".join(hdr)
        if "показатель" in joined and "как измеряется" in joined:
            return tbl
    return None


def collect_rows() -> list[dict]:
    rows: list[dict] = []
    for path in sorted(REG_DIR.glob("*.docx")):
        doc = Document(str(path))
        if not doc.tables:
            continue
        meta = _table_meta(doc.tables[0])
        reg_code = meta.get("Код регламента", "")
        pos_code = _position_code(meta, reg_code)
        pos_name = _position_name_ru(meta)
        dept = _dept_type(meta.get("Подразделение", ""), pos_code)
        kpi_tbl = _find_kpi_table(doc)
        if not kpi_tbl or len(kpi_tbl.rows) < 2:
            continue
        for i, row in enumerate(kpi_tbl.rows[1:], start=1):
            cells = [c.text.replace("\n", " ").strip() for c in row.cells]
            if len(cells) < 4:
                continue
            _no, name, how, target = cells[0], cells[1], cells[2], cells[3]
            if not name:
                continue
            code = f"KPI_{pos_code}_{i:02d}"
            formula = f"Как измеряется: {how}".strip()
            if target:
                formula = f"{formula} | Целевой ориентир: {target}"
            rows.append(
                {
                    "kpi_code": code,
                    "kpi_name": name[:256],
                    "unit": _infer_unit(target, how),
                    "period_type": "month",
                    "formula_or_rule": formula[:512] if len(formula) > 512 else formula,
                    "default_target": _infer_default_target(target),
                    "is_active": True,
                    "position_code": pos_code,
                    "primary_dept_type_code": dept,
                    "position_name_ru": pos_name[:256],
                }
            )
    return rows


def main() -> None:
    collected = collect_rows()
    df = pd.DataFrame(collected, columns=COLS)
    # Уникальность кодов (на случай дубликатов имён файлов)
    seen: set[str] = set()
    for i, c in enumerate(df["kpi_code"].tolist()):
        base = c
        n = 1
        while c in seen:
            n += 1
            c = f"{base}_DUP{n}"
        seen.add(c)
        df.at[i, "kpi_code"] = c

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    for dest in (OUT_PATH, OUT_DOWNLOADS):
        with pd.ExcelWriter(dest, engine="openpyxl") as w:
            df.to_excel(w, sheet_name=SHEET, index=False)
        print(f"Wrote {len(df)} rows to {dest}")


if __name__ == "__main__":
    main()
