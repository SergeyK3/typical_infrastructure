#!/usr/bin/env python3
"""
Синхронизация глобальных регламентов, KPI и навыков из DOCX (Google Drive / docs/regulations).

Источники:
  - docs/regulations/gdrive_default  — типовая enterprise-структура (default)
  - docs/regulations/gdrive_hosp     — медицинский контур стационара (hosp)

Локальные DOCX в корне docs/regulations/ не используются (не копируем сюда файлы с Drive).

Запуск из корня репозитория:
  python scripts/sync_global_regulations_from_sources.py
  python scripts/sync_global_regulations_from_sources.py --report reports/sync_report.json
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from uuid import NAMESPACE_URL, uuid5

from docx import Document
from openpyxl import load_workbook
from sqlalchemy import delete, select

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from app.db import SessionLocal
from app.models import (
    KpiTemplate,
    PositionDeptType,
    PositionRegulation,
    RegulationKpi,
)
from app.seed import _regulation_code_for_position
from skill_assessment.infrastructure.db_models import (
    CompetencyCatalogVersionRow,
    CompetencyMatrixRow,
    CompetencySkillDefinitionRow,
    KpiCatalogVersionRow,
    KpiDefinitionRow,
    KpiMatrixRow,
)

SECTION_RE = re.compile(r"^(\d+)\.\s+")
POSITION_IN_DUTY_RE = re.compile(r"\(([A-Z][A-Z0-9_]*)\)")
URL_RE = re.compile(r"https?://[^\s\)\]\"']+")

# DOCX position_code → position_code в каталоге шаблона hosp (legacy-коды в БД)
HOSP_POSITION_MAP: dict[str, str] = {
    "DOC_AMBUL": "ORDINATOR amb",
    "DOC_INPATIENT": "ORDINATOR hosp",
    "NURSE_AMBUL": "NURSE amb",
    "WARD_NURSE": "POST_NURSE",
    "NURSE_PROCEDURE": "PROCEDURE_NURSE",
    "CALL_REG": "CALL OPERATOR warm",
    "CALL_OUTBOUND": "CALL OPERATOR cold",
    "ADM_ZAM_AMBUL": "ADM_ZAM_POLYCLINIC",
    "ADM_ZAM_QUAL": "ADM_ZAM_QM",
}

DEFAULT_POSITION_MAP: dict[str, str] = {
    "DIRECTOR": "ADM_DIRECTOR",
    "SYSADMIN": "ADM_SYS_ADMIN",
}

# DOCX оставлены со старыми кодами должности после переименования в справочнике (все шаблоны).
RENAMED_POSITION_MAP: dict[str, str] = {
    "HR_MANAGER": "HR_GENERALIST",
    "ACC_MATERIAL_ACCOUNTANT": "ACC_ACCOUNTANT",
}

from scripts.fix_hosp_regulation_codes_and_urls import HOSP_REG_CANON

HOSP_MMC_DEPT: dict[str, str] = {x["position_code"]: x["dept_type_code"] for x in HOSP_REG_CANON}

MMC_POSITION_CODES = frozenset(
    {
        "DOC_AMBUL",
        "DOC_INPATIENT",
        "HEAD_DEPT",
        "ADM_ZAM_QUAL",
        "ADM_ZAMADM",
        "ADM_ZAM_AMBUL",
        "ADM_ZAM_LECH",
        "NURSE_AMBUL",
        "CALL_REG",
        "CALL_OUTBOUND",
        "WARD_NURSE",
        "NURSE_PROCEDURE",
        "ORDERLY",
        "NURSE_HOUSEKEEP",
        "DEPT_CHIEF_NURSE",
    }
)

REG_CODE_ALIASES: dict[str, str] = {
    "REG_DIRECTOR_V2": "REG_DIRECTOR_V1",
    "REG_HR_MANAGER_V2": "REG_HR_MANAGER_V1",
    "REG_HR_MANAGER_V1": "REG_HR_GENERALIST_V1",
    "REG_HR_RECRUITER_V2": "REG_HR_RECRUITER_V1",
    "REG_HR_HEAD_V2": "REG_HR_HEAD_V1",
    "REG_ACC_MATERIAL_ACCOUNTANT_V2": "REG_ACC_MATERIAL_ACCOUNTANT_V1",
    "REG_ACC_MATERIAL_ACCOUNTANT_V1": "REG_ACC_ACCOUNTANT_V1",
    "REG_ACC_CHIEF_ACCOUNTANT_V2": "REG_ACC_CHIEF_ACCOUNTANT_V1",
    "REG_SYSADMIN_V2": "REG_SYSADMIN_V1",
    "REG_MKT_MANAGER_V2": "REG_MKT_MANAGER_V1",
    "REG_SALES_MGR_V2": "REG_SALES_MGR_V1",
    "REG_SALES_TEAM_LEAD_V2": "REG_SALES_TEAM_LEAD_V1",
    "REG_QUAL_HEAD_V2": "REG_QUAL_HEAD_V1",
    "REG_QUAL_SPECIALIST_V2": "REG_QUAL_SPECIALIST_V1",
    "REG_LEADGEN_SPECIALIST_V2": "REG_LEADGEN_SPECIALIST_V1",
    "REG_INFO_SYSTEM_SUPPORT_V2": "REG_INFO_SYSTEM_SUPPORT_V1",
    "REG_PR_SPECIALIST_V2": "REG_PR_SPECIALIST_V1",
    "REG_PROD_TECH_DIR_V2": "REG_PROD_TECH_DIR_V1",
    "REG_ADM_ZAMADM_V2": "REG_ADM_ZAMADM_V1",
    # MMC-коды в DOCX → коды в seed default (если совпадают по должности)
    "REG_DOC_AMBUL_V1": "REG_DOC_AMBUL_V1",
    "REG_DOC_INPATIENT_V1": "REG_DOC_INPATIENT_V1",
}

WRONG_NAME_MARKERS = (
    "Руководитель отдела кадров",
    "Заместитель директора по производству (технический директор)",
    "Типовой шаблон",
)


@dataclass
class ParsedRegulation:
    source_file: str
    regulation_code: str
    position_code: str
    regulation_name: str | None = None
    goal_summary: str | None = None
    ckp_short: str | None = None
    ckp_full: str | None = None
    google_doc_url: str | None = None
    dept_type_code: str | None = None
    hard_skills: list[tuple[int, str]] = field(default_factory=list)
    soft_skills: list[tuple[int, str]] = field(default_factory=list)
    skills: list[tuple[int, str]] = field(default_factory=list)
    kpis: list[dict] = field(default_factory=list)


SOFT_SKILL_KIND = "soft job skill"
SOFT_SKILL_RANK_BASE = 8
HARD_SKILL_RANK_MAX = 7


@dataclass
class SyncReport:
    regulations_updated: list[dict] = field(default_factory=list)
    name_fixes: list[dict] = field(default_factory=list)
    ckp_short_derived: list[dict] = field(default_factory=list)
    kpi_templates_added: list[str] = field(default_factory=list)
    regulation_kpis_linked: list[str] = field(default_factory=list)
    skills_added: list[str] = field(default_factory=list)
    kpi_matrix_added: list[str] = field(default_factory=list)
    skipped_no_docx: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def _meta_table(doc: Document) -> dict[str, str]:
    meta: dict[str, str] = {}
    if not doc.tables:
        return meta
    for row in doc.tables[0].rows:
        cells = [c.text.replace("\n", " ").strip() for c in row.cells]
        if len(cells) >= 2 and cells[0]:
            meta[cells[0]] = cells[1]
    return meta


def _paragraphs(doc: Document) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def _sections(paras: list[str]) -> dict[int, list[str]]:
    sections: dict[int, list[str]] = {}
    current: int | None = None
    for line in paras:
        m = SECTION_RE.match(line)
        if m:
            current = int(m.group(1))
            sections.setdefault(current, []).append(line)
            continue
        if current is not None:
            sections[current].append(line)
    return sections


def _position_code(meta: dict[str, str], regulation_code: str) -> str:
    duty = meta.get("Должность", "")
    m = POSITION_IN_DUTY_RE.search(duty)
    if m:
        return m.group(1)
    m2 = re.match(r"REG_(.+)_V\d+$", regulation_code)
    return m2.group(1) if m2 else regulation_code


def _normalize_reg_code(code: str, position_code: str) -> str:
    c = (code or "").strip()
    for _ in range(4):
        if c in REG_CODE_ALIASES:
            c = REG_CODE_ALIASES[c]
        else:
            break
    if c.endswith("_V2"):
        c = c[:-3] + "_V1"
        for _ in range(4):
            if c in REG_CODE_ALIASES:
                c = REG_CODE_ALIASES[c]
            else:
                break
    if not c:
        return _regulation_code_for_position(position_code)
    return c


def _find_url(texts: list[str]) -> str | None:
    for t in texts:
        for m in URL_RE.finditer(t):
            return m.group(0)[:512]
    return None


def _derive_ckp_short(ckp_full: str | None, ckp_short: str | None) -> str | None:
    if ckp_short and ckp_short.strip():
        return ckp_short.strip()[:512]
    if not ckp_full:
        return None
    lines = [ln.strip() for ln in ckp_full.splitlines() if ln.strip()]
    candidate = None
    for line in lines:
        low = line.lower()
        if low.startswith("3.") or "ключевые результаты" in low:
            continue
        if len(line) < 15:
            continue
        candidate = line
        break
    if not candidate:
        return None
    # Первая фраза до точки с запятой или точки
    short = re.split(r"[.;]", candidate, maxsplit=1)[0].strip()
    # Глагол в неопределённой форме: «Пациенты получают» → «Получать»
    m = re.match(
        r"^(?:[А-ЯA-Z][а-яa-z]+(?:\s+[а-яa-z]+)*\s+)?(получа(?:ют|ет)|обеспечива(?:ют|ет)|"
        r"выполня(?:ют|ет)|осуществля(?:ют|ет)|провод(?:ят|ит)|формиру(?:ют|ет)|"
        r"поддержива(?:ют|ет)|организу(?:ют|ет)|контролиру(?:ют|ет)|снижа(?:ют|ет)|"
        r"своевременно\s+(\w+)|обеспечить|обеспечивать)\b",
        short,
        re.I,
    )
    if m:
        verb = m.group(1).lower()
        replacements = {
            "получают": "получать",
            "получает": "получать",
            "обеспечивают": "обеспечивать",
            "обеспечивает": "обеспечивать",
            "выполняют": "выполнять",
            "выполняет": "выполнять",
            "осуществляют": "осуществлять",
            "осуществляет": "осуществлять",
            "проводят": "проводить",
            "проводит": "проводить",
            "формируют": "формировать",
            "формирует": "формировать",
            "поддерживают": "поддерживать",
            "поддерживает": "поддерживать",
            "организуют": "organизовать",
            "организует": "организовать",
            "контролируют": "контролировать",
            "контролирует": "контролировать",
            "снижают": "снижать",
            "снижает": "снижать",
        }
        inf = replacements.get(verb, verb if verb.endswith("ть") else short)
        if inf != short and len(inf) < len(short):
            short = inf[0].upper() + inf[1:] if inf else short
    return short[:512]


def _parse_skills_table(tbl) -> list[tuple[int, str]]:
    if not tbl.rows:
        return []
    hdr = [c.text.strip().lower() for c in tbl.rows[0].cells]
    joined = " ".join(hdr)
    if "навык" not in joined or ("№" not in joined and "приоритет" not in joined):
        return []
    out: list[tuple[int, str]] = []
    for row in tbl.rows[1:]:
        cells = [c.text.replace("\n", " ").strip() for c in row.cells]
        if len(cells) < 2:
            continue
        rank_raw = cells[0]
        title = cells[1]
        if not title or title.lower().startswith("информация"):
            continue
        rank = int(rank_raw) if str(rank_raw).strip().isdigit() else len(out) + 1
        out.append((rank, title[:512]))
    return out


def _find_skill_tables(doc: Document) -> tuple[list[tuple[int, str]], list[tuple[int, str]]]:
    """Таблицы раздела 7: (hard skills, soft job skills).

    Одна таблица — только hard. Две таблицы — первая hard, вторая soft (шаблон ММЦ).
    """
    tables: list[list[tuple[int, str]]] = []
    for tbl in doc.tables:
        parsed = _parse_skills_table(tbl)
        if parsed:
            tables.append(parsed)
    if not tables:
        return [], []
    if len(tables) == 1:
        return tables[0], []
    return tables[0], tables[1]


def _find_skills_table(doc: Document) -> list[tuple[int, str]]:
    """Сводный список для отчётов: hard 1..7, soft 8..14 (если есть)."""
    hard, soft = _find_skill_tables(doc)
    if not hard and not soft:
        return []
    out = list(hard)
    rank = SOFT_SKILL_RANK_BASE
    for _table_rank, title in soft:
        out.append((rank, title))
        rank += 1
    return out


def _skill_code_part(value: str) -> str:
    return re.sub(r"[^\w]+", "_", (value or "").strip()).strip("_")[:32]


def _skill_code(
    template_code: str,
    position_code: str,
    rank: int,
    *,
    is_soft: bool,
) -> str:
    part = _skill_code_part(position_code)
    if is_soft:
        return f"CSOFT_{part}_{rank}"[:64]
    return f"C_{template_code}_{part}_{rank}"[:64]


def _find_kpi_table(doc: Document) -> list[dict]:
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        hdr = [c.text.strip().lower() for c in tbl.rows[0].cells]
        joined = " ".join(hdr)
        if "показатель" in joined and ("как измеряется" in joined or "код" in joined):
            out: list[dict] = []
            for i, row in enumerate(tbl.rows[1:], start=1):
                cells = [c.text.replace("\n", " ").strip() for c in row.cells]
                if len(cells) < 3:
                    continue
                name = cells[1] if len(cells) > 1 else ""
                if not name:
                    continue
                how = cells[2] if len(cells) > 2 else ""
                kpi_code_raw = cells[3] if len(cells) > 3 else ""
                target = cells[4] if len(cells) > 4 else ""
                if not target and len(cells) == 4 and not re.search(r"KPI_", kpi_code_raw or ""):
                    target = kpi_code_raw
                    kpi_code_raw = ""
                code_m = re.search(r"(KPI_[A-Z0-9_]+)", kpi_code_raw or "")
                kpi_code = code_m.group(1) if code_m else ""
                kpi_code = re.sub(r"\s+V\d+.*$", "", kpi_code).strip()
                out.append(
                    {
                        "kpi_code": kpi_code,
                        "kpi_name": name[:256],
                        "how": how[:256],
                        "target": target[:128],
                        "_row_index": i,
                    }
                )
            if out:
                return out
    return []


def parse_docx(path: Path, url_map: dict[str, str]) -> ParsedRegulation | None:
    doc = Document(str(path))
    meta = _meta_table(doc)
    raw_code = (meta.get("Код регламента") or "").strip()
    if not raw_code:
        return None
    position_code = _position_code(meta, raw_code)
    regulation_code = _normalize_reg_code(raw_code, position_code)
    paras = _paragraphs(doc)
    sec = _sections(paras)
    goal_parts = sec.get(2, [])
    goal_summary = " ".join(goal_parts[1:])[:512] if len(goal_parts) > 1 else None
    s3 = sec.get(3, [])
    ckp_short = (s3[1] if len(s3) > 1 else (s3[0] if s3 else None)) or None
    if ckp_short:
        ckp_short = ckp_short[:512]
    ckp_chunks: list[str] = []
    for n in (3, 4, 5):
        if n in sec:
            ckp_chunks.extend(sec[n])
    ckp_full = "\n\n".join(ckp_chunks) if ckp_chunks else None
    name = (meta.get("Должность") or "").replace("\n", " ")
    name = POSITION_IN_DUTY_RE.sub("", name).strip()
    name = re.sub(r"\s*\|\s*", " ", name).strip()
    regulation_name = (f"Регламент: {name}" if name else None)
    if regulation_name:
        regulation_name = regulation_name.replace("|", " ").strip()[:256]
    url = _find_url(paras) or url_map.get(path.stem) or url_map.get(path.name)
    subdivision = meta.get("Подразделение", "")
    dept = None
    if subdivision:
        dept = subdivision.split("/")[0].strip()[:32]
    hard_skills, soft_skills = _find_skill_tables(doc)
    return ParsedRegulation(
        source_file=str(path),
        regulation_code=regulation_code,
        position_code=position_code,
        regulation_name=regulation_name,
        goal_summary=goal_summary,
        ckp_short=ckp_short,
        ckp_full=ckp_full,
        google_doc_url=url,
        dept_type_code=dept,
        hard_skills=hard_skills,
        soft_skills=soft_skills,
        skills=_find_skills_table(doc),
        kpis=_find_kpi_table(doc),
    )


def load_url_maps() -> dict[str, str]:
    url_map: dict[str, str] = {}
    for folder in (ROOT / "docs/regulations/gdrive_default", ROOT / "docs/regulations/gdrive_hosp"):
        if not folder.is_dir():
            continue
        for xlsx in folder.glob("*.xlsx"):
            wb = load_workbook(xlsx, read_only=True, data_only=True)
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                rows = ws.iter_rows(values_only=True)
                header = next(rows, None)
                if not header:
                    continue
                hdr = [str(c).strip().lower() if c else "" for c in header]
                i_file = next((i for i, h in enumerate(hdr) if "файл" in h or h == "имя файла"), 0)
                i_url = next((i for i, h in enumerate(hdr) if "ссылка" in h or "url" in h), None)
                if i_url is None and len(header) >= 2:
                    i_url = 1 if "http" in str(header[1] or "") else (4 if len(header) > 4 else None)
                for row in rows:
                    if not row:
                        continue
                    fname = str(row[i_file] or "").strip()
                    if not fname:
                        continue
                    stem = Path(fname).stem
                    url = None
                    if i_url is not None and i_url < len(row) and row[i_url]:
                        url = str(row[i_url]).strip()
                    if not url:
                        for cell in row:
                            if cell and "docs.google.com" in str(cell):
                                url = str(cell).strip()
                                break
                    if url:
                        url_map[stem] = url[:512]
                        url_map[fname] = url[:512]
            wb.close()
    return url_map


def collect_parsed(url_map: dict[str, str]) -> dict[str, ParsedRegulation]:
    """Ключ — position_code из DOCX; при дубликатах побеждает gdrive_hosp над gdrive_default."""
    by_pos: dict[str, ParsedRegulation] = {}
    priority_dirs = [
        ROOT / "docs/regulations/gdrive_default",
        ROOT / "docs/regulations/gdrive_hosp",
    ]
    seen_files: set[str] = set()
    for folder in priority_dirs:
        if not folder.is_dir():
            continue
        for path in sorted(folder.glob("*.docx")):
            key = path.name.lower()
            if key in seen_files:
                continue
            seen_files.add(key)
            try:
                parsed = parse_docx(path, url_map)
            except Exception as exc:
                continue
            if parsed:
                by_pos[parsed.position_code] = parsed
    return by_pos


def _reg_kpi_id(template_code: str, regulation_code: str, kpi_code: str) -> str:
    return uuid5(NAMESPACE_URL, f"seed:reg_kpi:{template_code}:{regulation_code}:{kpi_code}").hex


def _target_position(template_code: str, docx_position: str) -> str:
    if template_code == "hosp" and docx_position in HOSP_POSITION_MAP:
        return HOSP_POSITION_MAP[docx_position]
    if template_code == "default" and docx_position in DEFAULT_POSITION_MAP:
        return DEFAULT_POSITION_MAP[docx_position]
    if docx_position in RENAMED_POSITION_MAP:
        return RENAMED_POSITION_MAP[docx_position]
    return docx_position


def _source_for_template(template_code: str, docx_position: str) -> bool:
    if template_code == "default":
        return docx_position not in MMC_POSITION_CODES or docx_position == "PR_SPECIALIST"
    if template_code == "hosp":
        return True
    return False


def _needs_name_fix(current: str | None, expected: str | None) -> bool:
    if not expected:
        return False
    cur = (current or "").strip()
    if not cur:
        return True
    for marker in WRONG_NAME_MARKERS:
        if marker in cur and marker not in expected:
            return True
    exp_core = expected.replace("Регламент:", "").strip().lower()
    cur_core = cur.replace("Регламент:", "").strip().lower()
    if exp_core and cur_core and exp_core not in cur_core and cur_core not in exp_core:
        # сильное расхождение
        if any(w in cur_core for w in ("кадров", "технический директор", "типовой шаблон")):
            return True
    return False


def _active_competency_version(db, template_code: str) -> CompetencyCatalogVersionRow | None:
    return db.scalar(
        select(CompetencyCatalogVersionRow)
        .where(
            CompetencyCatalogVersionRow.template_code == template_code,
            CompetencyCatalogVersionRow.status == "active",
        )
        .order_by(CompetencyCatalogVersionRow.published_at.desc())
        .limit(1)
    )


def _active_kpi_version(db, template_code: str) -> KpiCatalogVersionRow | None:
    del template_code  # одна глобальная KPI-версия на оба шаблона
    return db.scalar(
        select(KpiCatalogVersionRow)
        .where(
            KpiCatalogVersionRow.status == "active",
            KpiCatalogVersionRow.client_id.is_(None),
        )
        .order_by(KpiCatalogVersionRow.published_at.desc())
        .limit(1)
    )


def sync_template(
    db,
    template_code: str,
    by_pos: dict[str, ParsedRegulation],
    report: SyncReport,
) -> None:
    regulations = db.scalars(
        select(PositionRegulation).where(
            PositionRegulation.template_code == template_code,
            PositionRegulation.is_current == True,
        )
    ).all()

    comp_ver = _active_competency_version(db, template_code)
    kpi_ver = _active_kpi_version(db, template_code)

    for reg in regulations:
        pos = (reg.position_code or "").strip()
        # Найти DOCX по position_code или обратному маппингу
        parsed: ParsedRegulation | None = None
        for docx_pos, row in by_pos.items():
            if not _source_for_template(template_code, docx_pos):
                continue
            target = _target_position(template_code, docx_pos)
            if target == pos:
                parsed = row
                break
        if not parsed:
            if (not reg.ckp_short or not reg.ckp_full) or "Заполните" in (reg.goal_summary or ""):
                report.skipped_no_docx.append(f"{template_code}:{pos}")
            continue

        changed = False
        name_mismatch = (
            parsed.regulation_name
            and (reg.regulation_name or "").strip() != parsed.regulation_name.strip()
        )
        if _needs_name_fix(reg.regulation_name, parsed.regulation_name) or name_mismatch:
            if parsed.regulation_name and reg.regulation_name != parsed.regulation_name:
                report.name_fixes.append(
                    {
                        "template": template_code,
                        "position": pos,
                        "was": reg.regulation_name,
                        "now": parsed.regulation_name,
                    }
                )
                reg.regulation_name = parsed.regulation_name
                changed = True

        force_text = name_mismatch or "Заполните" in (reg.goal_summary or "")
        if parsed.goal_summary and (force_text or not reg.goal_summary):
            reg.goal_summary = parsed.goal_summary
            changed = True

        if parsed.ckp_full and (force_text or not (reg.ckp_full or "").strip()):
            reg.ckp_full = parsed.ckp_full
            changed = True

        new_short = _derive_ckp_short(reg.ckp_full or parsed.ckp_full, reg.ckp_short or parsed.ckp_short)
        if new_short and (force_text or not (reg.ckp_short or "").strip()):
            reg.ckp_short = new_short
            if not (reg.ckp_short or "").strip() == (parsed.ckp_short or "").strip():
                report.ckp_short_derived.append(f"{template_code}:{pos}")
            changed = True
        elif parsed.ckp_short and (force_text or not (reg.ckp_short or "").strip()):
            reg.ckp_short = parsed.ckp_short
            changed = True

        if parsed.google_doc_url and (
            not reg.google_doc_url
            or "example_" in (reg.google_doc_url or "")
            or "1BxNPYGYQtLED47zaC7iGv9D1" in (reg.google_doc_url or "")
            or "1jOgPLv6nYbwGcqJ2VfKeLbOC" in (reg.google_doc_url or "")
            or "1vNTwLOAroBUZFnwZDR33ZGs5" in (reg.google_doc_url or "")
        ):
            reg.google_doc_url = parsed.google_doc_url
            changed = True

        if changed:
            report.regulations_updated.append(
                {"template": template_code, "position": pos, "regulation_code": reg.regulation_code}
            )

        # KPI templates + links
        dept = db.scalar(
            select(PositionDeptType.dept_type_code).where(
                PositionDeptType.template_code == template_code,
                PositionDeptType.position_code == pos,
                PositionDeptType.is_primary == True,
            ).limit(1)
        ) or (parsed.dept_type_code or "ADM")

        if parsed.kpis:
            real_codes: list[str] = []
            for k in parsed.kpis:
                i = int(k.get("_row_index") or len(real_codes) + 1)
                base_code = (k.get("kpi_code") or "").strip()
                if not base_code or re.fullmatch(r"KPI_\d+", base_code):
                    base_code = f"KPI_{pos}_{i:02d}"
                kpi_code = base_code[:64]
                suffix = 1
                while db.get(KpiTemplate, (template_code, kpi_code)):
                    existing_kt = db.get(KpiTemplate, (template_code, kpi_code))
                    if existing_kt and existing_kt.position_code in (None, pos, parsed.position_code):
                        break
                    suffix += 1
                    kpi_code = f"{base_code}_{suffix}"[:64]

                existing_kt = db.get(KpiTemplate, (template_code, kpi_code))
                formula = k.get("how") or ""
                if k.get("target"):
                    formula = f"{formula} | Целевой ориентир: {k['target']}".strip(" |")
                if not existing_kt:
                    db.add(
                        KpiTemplate(
                            template_code=template_code,
                            kpi_code=kpi_code,
                            kpi_name=k["kpi_name"],
                            unit="%" if "%" in (k.get("target") or "") else "индекс",
                            period_type="month",
                            formula_or_rule=formula[:512] if formula else None,
                            default_target=None,
                            is_active=True,
                            position_code=pos,
                        )
                    )
                    report.kpi_templates_added.append(f"{template_code}:{kpi_code}")

                real_codes.append(kpi_code)

            # Удалить placeholder KPI_TMPL_* если есть реальные
            if real_codes:
                for rk in db.scalars(
                    select(RegulationKpi).where(
                        RegulationKpi.template_code == template_code,
                        RegulationKpi.regulation_code == reg.regulation_code,
                        RegulationKpi.kpi_code.like("KPI_TMPL_%"),
                    )
                ).all():
                    db.delete(rk)

            db.flush()
            existing_rk = {
                rk.kpi_code
                for rk in db.scalars(
                    select(RegulationKpi).where(
                        RegulationKpi.template_code == template_code,
                        RegulationKpi.regulation_code == reg.regulation_code,
                    )
                ).all()
            }
            for kpi_code in real_codes:
                if kpi_code in existing_rk:
                    continue
                db.add(
                    RegulationKpi(
                        id=_reg_kpi_id(template_code, reg.regulation_code, kpi_code),
                        template_code=template_code,
                        regulation_code=reg.regulation_code,
                        kpi_code=kpi_code,
                        target_value=None,
                        period_type="month",
                        weight=None,
                        is_required=True,
                    )
                )
                report.regulation_kpis_linked.append(f"{template_code}:{reg.regulation_code}:{kpi_code}")

        # Skills → active competency catalog
        skill_dept = HOSP_MMC_DEPT.get(pos, dept) if template_code == "hosp" else dept
        skills_to_sync: list[tuple[int, str, bool]] = []
        if template_code == "hosp" and parsed.soft_skills:
            # Soft job skills — точная копия из раздела Б; hard skills не трогаем (есть по умолчанию).
            for idx, (_table_rank, title) in enumerate(parsed.soft_skills):
                skills_to_sync.append((SOFT_SKILL_RANK_BASE + idx, title, True))
        elif parsed.hard_skills:
            for rank, title in parsed.hard_skills:
                skills_to_sync.append((rank, title, False))
        elif parsed.skills:
            for rank, title in parsed.skills:
                skills_to_sync.append((rank, title, False))

        if skills_to_sync and comp_ver:
            existing_ranks = {
                r.skill_rank
                for r in db.scalars(
                    select(CompetencyMatrixRow).where(
                        CompetencyMatrixRow.version_id == comp_ver.id,
                        CompetencyMatrixRow.position_code == pos,
                        CompetencyMatrixRow.department_code == skill_dept,
                    )
                ).all()
            }
            for rank, title, is_soft in skills_to_sync:
                if rank in existing_ranks:
                    continue
                sid = str(uuid.uuid4())
                scode = _skill_code(template_code, pos, rank, is_soft=is_soft)
                db.add(
                    CompetencySkillDefinitionRow(
                        id=sid,
                        client_id=None,
                        template_code=template_code,
                        skill_code=scode,
                        title_ru=title[:512],
                        description=SOFT_SKILL_KIND if is_soft else None,
                        is_active=True,
                    )
                )
                db.add(
                    CompetencyMatrixRow(
                        id=str(uuid.uuid4()),
                        version_id=comp_ver.id,
                        position_code=pos,
                        department_code=skill_dept,
                        skill_definition_id=sid,
                        skill_rank=rank,
                        is_active=True,
                    )
                )
                kind = "soft" if is_soft else "hard"
                report.skills_added.append(f"{template_code}:{pos}:{rank}:{kind}")
                existing_ranks.add(rank)

        # KPI matrix in skill_assessment
        if parsed.kpis and kpi_ver:
            db.flush()
            existing_kranks = {
                r.kpi_rank
                for r in db.scalars(
                    select(KpiMatrixRow).where(
                        KpiMatrixRow.version_id == kpi_ver.id,
                        KpiMatrixRow.position_code == pos,
                        KpiMatrixRow.department_code == dept,
                    )
                ).all()
            }
            for rank, k in enumerate(parsed.kpis, start=1):
                if rank in existing_kranks:
                    continue
                i = int(k.get("_row_index") or rank)
                kpi_code = (k.get("kpi_code") or "").strip()
                if not kpi_code or re.fullmatch(r"KPI_\d+", kpi_code):
                    kpi_code = f"KPI_{pos}_{i:02d}"
                kid = db.scalar(
                    select(KpiDefinitionRow.id).where(
                        KpiDefinitionRow.kpi_code == kpi_code,
                        KpiDefinitionRow.client_id.is_(None),
                    )
                )
                if not kid:
                    kid = str(uuid.uuid4())
                    db.add(
                        KpiDefinitionRow(
                            id=kid,
                            client_id=None,
                            kpi_code=kpi_code,
                            title_ru=k["kpi_name"][:512],
                            unit="%" if "%" in (k.get("target") or "") else "индекс",
                            period_type="month",
                            default_target=None,
                            is_active=True,
                        )
                    )
                    db.flush()
                db.add(
                    KpiMatrixRow(
                        id=str(uuid.uuid4()),
                        version_id=kpi_ver.id,
                        position_code=pos,
                        department_code=dept,
                        kpi_definition_id=kid,
                        kpi_rank=rank,
                        is_active=True,
                    )
                )
                existing_kranks.add(rank)
                report.kpi_matrix_added.append(f"{template_code}:{pos}:{kpi_code}")


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--report", type=str, default="", help="JSON-отчёт")
    args = parser.parse_args()

    url_map = load_url_maps()
    by_pos = collect_parsed(url_map)
    report = SyncReport()

    db = SessionLocal()
    try:
        for tpl in ("default", "hosp"):
            sync_template(db, tpl, by_pos, report)
        db.commit()
    finally:
        db.close()

    # Консольный отчёт
    print(f"DOCX parsed (unique positions): {len(by_pos)}")
    print(f"Regulations updated: {len(report.regulations_updated)}")
    print(f"Name fixes: {len(report.name_fixes)}")
    print(f"CKP short derived: {len(report.ckp_short_derived)}")
    print(f"KPI templates added: {len(report.kpi_templates_added)}")
    print(f"Regulation-KPI links added: {len(report.regulation_kpis_linked)}")
    print(f"Skills added: {len(report.skills_added)}")
    print(f"KPI matrix rows added: {len(report.kpi_matrix_added)}")
    print(f"Skipped (no DOCX): {len(report.skipped_no_docx)}")
    if report.name_fixes:
        print("\n--- Name fixes ---")
        for item in report.name_fixes[:20]:
            print(f"  {item['template']} {item['position']}: {item['was']!r} -> {item['now']!r}")
    if report.skipped_no_docx:
        print("\n--- No DOCX ---")
        for s in report.skipped_no_docx:
            print(f"  {s}")
    if report.warnings:
        print("\n--- Warnings ---")
        for w in report.warnings:
            print(f"  {w}")

    if args.report:
        out = Path(args.report)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(json.dumps(report.__dict__, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"\nReport written to {out}")


if __name__ == "__main__":
    main()
