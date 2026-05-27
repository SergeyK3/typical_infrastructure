#!/usr/bin/env python3
"""
Дописать раздел «7. Ключевые навыки» (hard & soft job skills) в существующие Google Docs.

Сохраняет file_id и URL — экспорт → патч DOCX → upload в тот же документ.

Папки:
  medical: 1BAKymXUdfTO_UKeXZHDbczGV5TOml88-  (ММЦ, A/B таблицы)
  admin:   1JzLlrGCeut77aXmmEL9pcs7QP1RcGuBq  (enterprise, hard skills из top20)

Запуск:
  python scripts/push_mmc_regulations_skills_to_gdrive.py
  python scripts/push_mmc_regulations_skills_to_gdrive.py --folder medical
  python scripts/push_mmc_regulations_skills_to_gdrive.py --dry-run
"""

from __future__ import annotations

import argparse
import io
import json
import re
import sys
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from dotenv import load_dotenv
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
load_dotenv(ROOT / ".env")

from psychological_testing.integration.google_drive_client import _build_drive_service  # noqa: E402
from scripts.mmc_docx_common import (  # noqa: E402
    SKILLS_SECTION_INTRO,
    SKILLS_SECTION_INTRO2,
    fill_skills_table,
    format_regulation_tables,
    has_mmc_ab_skills_section,
    insert_mmc_skills_before_kpi,
    insert_paragraph_before,
    insert_table_before,
    position_code_from_filename,
    renumber_sections_from,
    find_kpi_heading,
)
from skill_assessment.data.top20_position_skills import TOP20_POSITION_SKILL_ROWS  # noqa: E402

MEDICAL_FOLDER_ID = "1BAKymXUdfTO_UKeXZHDbczGV5TOml88-"
ADMIN_FOLDER_ID = "1JzLlrGCeut77aXmmEL9pcs7QP1RcGuBq"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
GOOGLE_DOC_MIME = "application/vnd.google-apps.document"

ADMIN_POSITION_ALIASES = {
    "DIRECTOR": "ADM_DIRECTOR",
    "SYSADMIN": "ADM_SYS_ADMIN",
    "SALES_MGR": "SALES_MANAGER",
    "HR_MANAGER": "HR_GENERALIST",
    "ACC_MATERIAL_ACCOUNTANT": "ACC_ACCOUNTANT",
}

ADMIN_SKILLS_BY_POSITION = {
    code: skills for code, _name, skills in TOP20_POSITION_SKILL_ROWS
}

SKIP_NAME_MARKERS = (
    "симулятор экзамена",
    "каталог",
    "перечень регламентов",
)


@dataclass
class PushReport:
    updated: list[dict] = field(default_factory=list)
    skipped: list[str] = field(default_factory=list)
    failed: list[dict] = field(default_factory=list)


def _position_code_from_meta(doc: Document) -> str | None:
    if not doc.tables:
        return None
    meta = doc.tables[0]
    duty = meta.cell(2, 1).text.replace("\n", " ").strip()
    m = re.search(r"\(([A-Z][A-Z0-9_]*)\)", duty)
    if m:
        code = m.group(1)
    else:
        regulation_code = meta.cell(0, 1).text.strip()
        m2 = re.match(r"REG_(.+)_V\d+$", regulation_code)
        code = m2.group(1) if m2 else None
    if not code:
        return None
    return ADMIN_POSITION_ALIASES.get(code, code)


def has_admin_skills_section(doc: Document) -> bool:
    for para in doc.paragraphs:
        t = para.text.strip()
        if t.startswith("7.") and "навыки" in t.lower():
            return True
    return False


def insert_admin_skills_before_kpi(doc: Document, skills: list[str]) -> bool:
    if has_admin_skills_section(doc):
        return False
    kpi_heading = find_kpi_heading(doc)
    renumber_sections_from(doc, 7)
    insert_paragraph_before(kpi_heading, "7. Ключевые навыки")
    insert_paragraph_before(kpi_heading, SKILLS_SECTION_INTRO)
    skills_table = insert_table_before(doc, kpi_heading, rows=1, cols=4)
    fill_skills_table(skills_table, skills)
    insert_paragraph_before(kpi_heading, SKILLS_SECTION_INTRO2)
    format_regulation_tables(doc, [skills_table])
    return True


def list_google_docs(service, folder_id: str) -> list[dict]:
    q = (
        f"'{folder_id}' in parents and trashed = false "
        f"and mimeType = '{GOOGLE_DOC_MIME}'"
    )
    items: list[dict] = []
    page_token = None
    while True:
        resp = (
            service.files()
            .list(
                q=q,
                fields="nextPageToken, files(id,name,capabilities)",
                pageSize=100,
                pageToken=page_token,
                supportsAllDrives=True,
                includeItemsFromAllDrives=True,
            )
            .execute()
        )
        items.extend(resp.get("files", []))
        page_token = resp.get("nextPageToken")
        if not page_token:
            break
    return items


def export_docx(service, file_id: str) -> bytes:
    request = service.files().export_media(fileId=file_id, mimeType=DOCX_MIME)
    buf = io.BytesIO()
    downloader = MediaIoBaseDownload(buf, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    return buf.getvalue()


def upload_docx_inplace(service, file_id: str, docx_bytes: bytes) -> dict:
    media = MediaIoBaseUpload(io.BytesIO(docx_bytes), mimetype=DOCX_MIME, resumable=False)
    return (
        service.files()
        .update(
            fileId=file_id,
            media_body=media,
            supportsAllDrives=True,
            fields="id,modifiedTime,webViewLink",
        )
        .execute()
    )


def patch_medical_doc(doc: Document, name: str) -> tuple[bool, str]:
    position_code = position_code_from_filename(name)
    if not position_code:
        return False, "no_mmc_position_code"
    if has_mmc_ab_skills_section(doc):
        return False, "skills_ready"
    if not insert_mmc_skills_before_kpi(doc, position_code):
        return False, "patch_not_applied"
    return True, position_code


def patch_admin_doc(doc: Document, name: str) -> tuple[bool, str]:
    position_code = _position_code_from_meta(doc) or position_code_from_filename(name)
    if not position_code:
        return False, "no_position_code"
    position_code = ADMIN_POSITION_ALIASES.get(position_code, position_code)
    skills = ADMIN_SKILLS_BY_POSITION.get(position_code)
    if not skills:
        return False, f"no_skills_for:{position_code}"
    if has_admin_skills_section(doc):
        return False, "skills_ready"
    if not insert_admin_skills_before_kpi(doc, skills):
        return False, "patch_not_applied"
    return True, position_code


def should_skip_name(name: str) -> bool:
    low = name.lower()
    return any(m in low for m in SKIP_NAME_MARKERS)


def process_folder(
    service,
    folder_id: str,
    folder_kind: str,
    report: PushReport,
    *,
    dry_run: bool,
) -> None:
    for item in sorted(list_google_docs(service, folder_id), key=lambda x: x["name"]):
        name = item["name"]
        file_id = item["id"]
        if should_skip_name(name):
            report.skipped.append(f"{name}:non_regulation")
            continue
        if not item.get("capabilities", {}).get("canEdit"):
            report.failed.append({"name": name, "file_id": file_id, "reason": "canEdit=false"})
            continue

        try:
            docx_bytes = export_docx(service, file_id)
            tmp = Path(tempfile.gettempdir()) / f"gdrive_patch_{file_id}.docx"
            tmp.write_bytes(docx_bytes)
            doc = Document(str(tmp))

            if folder_kind == "medical":
                changed, reason = patch_medical_doc(doc, name)
            else:
                changed, reason = patch_admin_doc(doc, name)

            if not changed:
                report.skipped.append(f"{name}:{reason}")
                continue

            if dry_run:
                report.updated.append({"name": name, "file_id": file_id, "dry_run": True, "reason": reason})
                continue

            doc.save(str(tmp))
            meta = upload_docx_inplace(service, file_id, tmp.read_bytes())
            report.updated.append(
                {
                    "name": name,
                    "file_id": file_id,
                    "position": reason,
                    "modifiedTime": meta.get("modifiedTime"),
                    "url": f"https://docs.google.com/document/d/{file_id}/edit",
                }
            )
        except Exception as exc:
            report.failed.append({"name": name, "file_id": file_id, "reason": str(exc)})


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--folder",
        choices=("medical", "admin", "both"),
        default="both",
    )
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--report", default="")
    args = parser.parse_args()

    service = _build_drive_service()
    report = PushReport()

    folders: list[tuple[str, str]] = []
    if args.folder in ("medical", "both"):
        folders.append(("medical", MEDICAL_FOLDER_ID))
    if args.folder in ("admin", "both"):
        folders.append(("admin", ADMIN_FOLDER_ID))

    for kind, fid in folders:
        process_folder(service, fid, kind, report, dry_run=args.dry_run)

    print(f"Updated: {len(report.updated)}")
    print(f"Skipped: {len(report.skipped)}")
    print(f"Failed: {len(report.failed)}")

    if report.updated:
        print("\n--- Updated ---")
        for item in report.updated:
            print(f"  {item['name']} ({item.get('position', item.get('reason', ''))})")

    if report.skipped:
        print("\n--- Skipped ---")
        for item in report.skipped:
            print(f"  {item}")

    if report.failed:
        print("\n--- Failed ---")
        for item in report.failed:
            print(f"  {item['name']}: {item['reason']}")

    if args.report:
        out = Path(args.report)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(json.dumps(report.__dict__, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"\nReport: {out}")

    if report.failed and not args.dry_run:
        sys.exit(1)


if __name__ == "__main__":
    main()
