#!/usr/bin/env python3
"""
Добавить раздел «7. Ключевые навыки» в медицинские регламенты ММЦ (26.05.2026),
сдвинуть нумерацию разделов 7→8 … 12→13.

Обрабатывает docs/regulations/*.docx и docs/regulations/gdrive_hosp/*.docx
(файлы из SPECS build_mmc_regulations_docx.py).

Запуск из корня репозитория:
  python scripts/patch_mmc_regulations_skills_section.py
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from scripts.mmc_docx_common import (  # noqa: E402
    SKILLS_BY_CODE,
    format_regulation_tables,
    setup_section7_skills,
)

MMC_FILENAMES = [
    "Регламент_Заместитель_директора_по_медицинской_части_ADM_ZAM_LECH.docx",
    "Регламент_Заместитель_директора_по_административным_вопросам_ADM_ZAMADM.docx",
    "Регламент_Заместитель_директора_по_амбулаторно-поликлинической_работе_ADM_ZAM_AMBUL.docx",
    "Регламент_Заместитель_главного_врача_по_экспертизе_качества_ADM_ZAM_QUAL.docx",
    "Регламент_Врач_амбулаторного_приёма_DOC_AMBUL.docx",
    "Регламент_Врач_госпитального_отделения_DOC_INPATIENT.docx",
    "Регламент_Заведующий_госпитальным_отделением_HEAD_DEPT.docx",
    "Регламент_Медсестра_амбулаторного_приёма_NURSE_AMBUL.docx",
    "Регламент_Постовая_медсестра_WARD_NURSE.docx",
    "Регламент_Старшая_медсестра_госпитального_отделения_DEPT_CHIEF_NURSE.docx",
    "Регламент_Оператор_колл-центра_регистрация_CALL_REG.docx",
    "Регламент_Оператор_колл-центра_холодные_звонки_CALL_OUTBOUND.docx",
    "Регламент_Процедурная_медсестра_NURSE_PROCEDURE.docx",
    "Регламент_Сестра-хозяйка_NURSE_HOUSEKEEP.docx",
    "Регламент_Санитарка_ORDERLY.docx",
]

def _position_code_from_path(path: Path) -> str | None:
    stem = path.stem
    for code in sorted(SKILLS_BY_CODE, key=len, reverse=True):
        if stem.endswith(f"_{code}"):
            return code
    return None


def _has_ab_skills_subsections(doc: Document) -> bool:
    has_a = has_b = False
    for para in doc.paragraphs:
        t = para.text.strip().lower()
        if t.startswith("а.") and "hard" in t:
            has_a = True
        if t.startswith("б.") and "soft" in t:
            has_b = True
    return has_a and has_b


def patch_docx(path: Path) -> bool:
    code = _position_code_from_path(path)
    if not code:
        return False
    doc = Document(str(path))
    setup_section7_skills(doc, code)
    format_regulation_tables(doc)
    doc.save(str(path))
    return True


def main() -> None:
    targets: list[Path] = []
    for folder in (ROOT / "docs/regulations", ROOT / "docs/regulations/gdrive_hosp"):
        for fname in MMC_FILENAMES:
            p = folder / fname
            if p.is_file():
                targets.append(p)

    patched = 0
    for path in targets:
        try:
            if _has_ab_skills_subsections(Document(str(path))):
                print(f"SKIP (A/B skills ready) {path.relative_to(ROOT)}")
                continue
            if not _position_code_from_path(path):
                print(f"SKIP (unknown position code) {path.relative_to(ROOT)}")
                continue
            if patch_docx(path):
                print(f"OK {path.relative_to(ROOT)}")
                patched += 1
        except Exception as exc:
            print(f"FAIL {path}: {exc}", file=sys.stderr)
            raise

    print(f"\nPatched {patched} file(s).")


if __name__ == "__main__":
    main()
