"""Генерация текстового регламента (DOCX) по названию должности с опорой на интернет и шаблон."""

from __future__ import annotations

import json
import re
import shutil
import uuid
from dataclasses import dataclass, field
from datetime import date
from html import unescape
from pathlib import Path

import httpx
from docx import Document

ROOT = Path(__file__).resolve().parents[2]
DOC_DIR = ROOT / "docs" / "regulations"
GENERATED_DIR = DOC_DIR / "generated"
TEMPLATE_CANDIDATES = (
    DOC_DIR / "Регламент_Специалист_по_связям_с_общественностью_PR_SPECIALIST.docx",
    DOC_DIR / "gdrive_default" / "Регламент_HR_менеджер_HR_MANAGER.docx",
)

SKIP_HEADINGS = (
    "РЕГЛАМЕНТ",
    "1.",
    "2.",
    "3.",
    "4.",
    "5.",
    "6.",
    "7.",
    "8.",
    "9.",
    "10.",
    "11.",
    "12.",
    "Согласование",
)

CYRILLIC_TO_LAT = {
    "а": "a",
    "б": "b",
    "в": "v",
    "г": "g",
    "д": "d",
    "е": "e",
    "ё": "e",
    "ж": "zh",
    "з": "z",
    "и": "i",
    "й": "y",
    "к": "k",
    "л": "l",
    "м": "m",
    "н": "n",
    "о": "o",
    "п": "p",
    "р": "r",
    "с": "s",
    "т": "t",
    "у": "u",
    "ф": "f",
    "х": "h",
    "ц": "ts",
    "ч": "ch",
    "ш": "sh",
    "щ": "sch",
    "ъ": "",
    "ы": "y",
    "ь": "",
    "э": "e",
    "ю": "yu",
    "я": "ya",
}


@dataclass
class RegulationDraft:
    position_title: str
    position_code: str
    regulation_code: str
    regulation_name: str
    goal_summary: str
    ckp_short: str
    ckp_full: str
    content_paragraphs: list[str] = field(default_factory=list)
    cycle_rows: list[list[str]] = field(default_factory=list)
    kpi_rows: list[list[str]] = field(default_factory=list)
    skill_rows: list[list[str]] = field(default_factory=list)
    meta: dict[str, str] = field(default_factory=dict)
    sources: list[str] = field(default_factory=list)
    notes: str | None = None


def resolve_template_docx() -> Path:
    for path in TEMPLATE_CANDIDATES:
        if path.is_file():
            return path
    gdrive = sorted((DOC_DIR / "gdrive_default").glob("*.docx"))
    if gdrive:
        return gdrive[0]
    local = sorted(DOC_DIR.glob("*.docx"))
    if local:
        return local[0]
    raise FileNotFoundError("regulation_template_docx_not_found")


def slug_position_code(title: str) -> str:
    raw = title.strip().lower()
    out: list[str] = []
    for ch in raw:
        if ch in CYRILLIC_TO_LAT:
            out.append(CYRILLIC_TO_LAT[ch])
        elif ch.isascii() and ch.isalnum():
            out.append(ch)
        elif ch in " -_/.":
            out.append("_")
    code = re.sub(r"_+", "_", "".join(out)).strip("_").upper()
    if not code:
        code = "NEW_POSITION"
    return code[:48]


def _set_cell(cell, text: str) -> None:
    cell.text = text


def _set_para(para, text: str) -> None:
    if not para.runs:
        para.text = text
        return
    para.runs[0].text = text
    for run in para.runs[1:]:
        run.text = ""


def _fill_table(table, rows: list[list[str]]) -> None:
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    for ri, row_data in enumerate(rows):
        while len(table.rows) <= ri:
            table.add_row()
        row = table.rows[ri]
        for ci, value in enumerate(row_data):
            while len(row.cells) <= ci:
                row.add_cell()
            _set_cell(row.cells[ci], value)


def analyze_template_structure(template_path: Path) -> dict:
    doc = Document(str(template_path))
    headings = [p.text.strip() for p in doc.paragraphs if re.match(r"^\d+\.\s+", p.text.strip())]
    fillable = 0
    for pi, para in enumerate(doc.paragraphs):
        if pi == 1:
            continue
        text = para.text.strip()
        if not text or text.startswith(SKIP_HEADINGS):
            continue
        fillable += 1
    return {
        "template_path": str(template_path),
        "headings": headings[:12],
        "fillable_paragraphs": fillable,
        "tables": len(doc.tables),
    }


def search_web_snippets(query: str, *, max_results: int = 6) -> list[dict[str, str]]:
    results: list[dict[str, str]] = []
    try:
        with httpx.Client(timeout=35.0, follow_redirects=True) as client:
            resp = client.post(
                "https://html.duckduckgo.com/html/",
                data={"q": query},
                headers={"User-Agent": "TypicalInfrastructure/1.0"},
            )
            resp.raise_for_status()
            html = resp.text
    except Exception:
        return results

    for block in re.findall(
        r'class="result__a"[^>]*href="([^"]+)"[^>]*>(.*?)</a>.*?class="result__snippet"[^>]*>(.*?)</',
        html,
        flags=re.S,
    ):
        url, title_raw, snippet_raw = block
        title = unescape(re.sub(r"<[^>]+>", "", title_raw)).strip()
        snippet = unescape(re.sub(r"<[^>]+>", "", snippet_raw)).strip()
        if not snippet:
            continue
        results.append({"title": title, "url": url, "snippet": snippet})
        if len(results) >= max_results:
            break
    return results


def _llm_compose(position_title: str, comment: str | None, snippets: list[dict[str, str]]) -> dict | None:
    try:
        from psychological_testing.services.llm_service import default_llm_model, get_llm_client
    except ImportError:
        return None

    sources_text = "\n".join(
        f"- {s.get('title', '')}: {s.get('snippet', '')} ({s.get('url', '')})" for s in snippets[:6]
    )
    if not sources_text.strip():
        sources_text = "(источники не найдены — используй типовую структуру должностного регламента)"

    system = (
        "Ты составляешь черновик должностного регламента на русском языке для типовой организации. "
        "Ответ — только JSON без markdown."
    )
    user = f"""Должность: {position_title}
Дополнительный комментарий заказчика: {comment or "—"}

Материалы из интернета:
{sources_text}

Верни JSON с ключами:
- goal_summary (строка, до 400 символов)
- ckp_short (строка, до 200 символов, глагол в неопределённой форме)
- ckp_results (массив из 4 строк — ключевые результаты)
- responsibilities (массив из 5 строк — зоны ответственности)
- intro (массив из 3 строк — общие положения)
- powers (массив из 3 строк — полномочия)
- accountability (массив из 2 строк — ответственность)
- qualifications (массив из 3 строк — квалификация)
- competencies (массив из 3 строк — компетенции)
- skills (массив из 7 строк — ключевые навыки)
- kpis (массив объектов {{name, measure, target}} — 5 штук)
- cycle (массив объектов {{period, actions, result}} — 3 штуги: Ежедневно, Еженедельно, Ежемесячно)
"""
    try:
        llm = get_llm_client()
        raw = llm.chat(default_llm_model("report"), [{"role": "system", "content": system}, {"role": "user", "content": user}], temperature=0.3)
        raw = raw.strip()
        if raw.startswith("```"):
            raw = re.sub(r"^```(?:json)?\s*", "", raw)
            raw = re.sub(r"\s*```$", "", raw)
        data = json.loads(raw)
        return data if isinstance(data, dict) else None
    except Exception:
        return None


def _fallback_compose(position_title: str, comment: str | None, snippets: list[dict[str, str]]) -> dict:
    blob = " ".join(s.get("snippet", "") for s in snippets).strip()
    if comment:
        blob = f"{comment}. {blob}".strip()
    if len(blob) < 80:
        blob = (
            f"{position_title} обеспечивает выполнение задач в зоне своей ответственности, "
            "соблюдает стандарты компании, своевременно эскалирует риски и поддерживает "
            "прозрачную отчётность для руководства."
        )

    def clip(text: str, n: int) -> str:
        t = re.sub(r"\s+", " ", text).strip()
        return t if len(t) <= n else t[: n - 1].rsplit(" ", 1)[0] + "…"

    goal = clip(f"Обеспечить эффективное выполнение функций должности «{position_title}» в соответствии со стандартами компании.", 400)
    ckp_short = clip(f"Обеспечивать результат по профилю должности «{position_title}» без критичных срывов сроков и качества.", 200)

    def paras(prefix: str, n: int) -> list[str]:
        base = clip(blob, 220)
        return [clip(f"{prefix}: {base}", 320) for _ in range(n)]

    return {
        "goal_summary": goal,
        "ckp_short": ckp_short,
        "intro": paras(f"Сотрудник на должности «{position_title}»", 3),
        "ckp_results": paras("Ключевой результат", 4),
        "responsibilities": paras("Зона ответственности", 5),
        "powers": paras("Полномочие", 3),
        "accountability": paras("Подотчётность", 2),
        "qualifications": paras("Требование", 3),
        "competencies": paras("Компетенция", 3),
        "skills": [clip(f"Навык {i + 1}: {blob[:120]}", 200) for i in range(7)],
        "kpis": [
            {"name": "Выполнение плана работ", "measure": "План-факт за период", "target": "100%"},
            {"name": "Соблюдение сроков", "measure": "Доля задач в SLA", "target": "≥ 95%"},
            {"name": "Качество результата", "measure": "Аудиты / проверки", "target": "без критичных замечаний"},
            {"name": "Дисциплина отчётности", "measure": "Своевременность отчётов", "target": "≥ 95%"},
            {"name": "Эскалация рисков", "measure": "Своевременность эскалаций", "target": "100%"},
        ],
        "cycle": [
            {
                "period": "Ежедневно",
                "actions": f"Оперативная работа по профилю «{position_title}», контроль сроков и качества.",
                "result": "Задачи дня выполнены, риски зафиксированы.",
            },
            {
                "period": "Еженедельно",
                "actions": "Сводка результатов, разбор отклонений, согласование приоритетов со смежными службами.",
                "result": "Неделя закрыта предсказуемо, отклонения устраняются.",
            },
            {
                "period": "Ежемесячно",
                "actions": "Анализ KPI, предложения по улучшению процессов, отчётность руководству.",
                "result": "Управляемая система работы и план улучшений.",
            },
        ],
    }


def compose_draft(position_title: str, comment: str | None, snippets: list[dict[str, str]]) -> RegulationDraft:
    title = position_title.strip()
    position_code = slug_position_code(title)
    regulation_code = f"REG_{position_code}_V1"
    data = _llm_compose(title, comment, snippets) or _fallback_compose(title, comment, snippets)

    ckp_results = [str(x).strip() for x in data.get("ckp_results") or [] if str(x).strip()]
    ckp_full_parts = ["3. Ключевые результаты должности", ""] + ckp_results
    ckp_full = "\n\n".join(ckp_full_parts)

    intro = [str(x).strip() for x in data.get("intro") or [] if str(x).strip()]
    responsibilities = [str(x).strip() for x in data.get("responsibilities") or [] if str(x).strip()]
    powers = [str(x).strip() for x in data.get("powers") or [] if str(x).strip()]
    accountability = [str(x).strip() for x in data.get("accountability") or [] if str(x).strip()]
    qualifications = [str(x).strip() for x in data.get("qualifications") or [] if str(x).strip()]
    competencies = [str(x).strip() for x in data.get("competencies") or [] if str(x).strip()]

    content: list[str] = []
    content.extend(intro[:3])
    content.append(str(data.get("goal_summary") or "").strip())
    content.extend(ckp_results[:4])
    content.extend(responsibilities[:6])
    content.extend(powers[:3])
    content.extend(accountability[:2])
    content.extend(qualifications[:3])
    content.extend(competencies[:3])
    content.append(
        f"Ценный конечный продукт должности «{title}»: "
        f"{data.get('ckp_short') or (ckp_results[0] if ckp_results else f'Обеспечить результат по должности «{title}».')}"
    )

    skills = [str(x).strip() for x in data.get("skills") or [] if str(x).strip()]
    skill_rows = [["№", "Навык", "Приоритет", "Версия"]]
    for i, skill in enumerate(skills[:7], start=1):
        skill_rows.append([str(i), skill, str(i), "V1 черновик"])

    kpis = data.get("kpis") or []
    kpi_rows = [["№", "Показатель", "Как измеряется", "Целевой ориентир"]]
    for i, row in enumerate(kpis[:5], start=1):
        if isinstance(row, dict):
            kpi_rows.append(
                [
                    str(i),
                    str(row.get("name") or f"KPI {i}"),
                    str(row.get("measure") or "План-факт"),
                    str(row.get("target") or "100%"),
                ]
            )

    cycle = data.get("cycle") or []
    cycle_rows = [["Период", "Основные действия", "Результат"]]
    for row in cycle[:3]:
        if isinstance(row, dict):
            cycle_rows.append(
                [
                    str(row.get("period") or ""),
                    str(row.get("actions") or ""),
                    str(row.get("result") or ""),
                ]
            )

    today = date.today().strftime("%d.%m.%Y")
    meta = {
        "Код регламента": regulation_code,
        "Версия": "V1",
        "Подразделение": "— уточнить —",
        "Статус": "Черновик",
        "Должность": f"{title}\n({position_code})",
        "Дата": today,
        "Область применения": comment.strip()[:240] if comment and comment.strip() else f"Типовой регламент для должности «{title}»",
        "Основание": "Сгенерировано из открытых источников; требует проверки HR/руководителя",
    }

    sources = [s.get("url", "") for s in snippets if s.get("url")]
    notes = None
    if sources:
        notes = "Источники при генерации: " + "; ".join(sources[:5])

    return RegulationDraft(
        position_title=title,
        position_code=position_code,
        regulation_code=regulation_code,
        regulation_name=f"Регламент: {title}",
        goal_summary=str(data.get("goal_summary") or "").strip(),
        ckp_short=str(data.get("ckp_short") or "").strip(),
        ckp_full=ckp_full,
        content_paragraphs=content,
        cycle_rows=cycle_rows,
        kpi_rows=kpi_rows,
        skill_rows=skill_rows,
        meta=meta,
        sources=sources,
        notes=notes,
    )


def _table_header_text(table) -> str:
    if not table.rows:
        return ""
    return " ".join(c.text.strip().lower() for c in table.rows[0].cells)


def _find_table_by_header(doc: Document, *keywords: str):
    for table in doc.tables:
        hdr = _table_header_text(table)
        if any(k in hdr for k in keywords):
            return table
    return None


def render_docx(draft: RegulationDraft, output_path: Path, *, template_path: Path | None = None) -> Path:
    template = template_path or resolve_template_docx()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, output_path)
    doc = Document(str(output_path))

    for row in doc.tables[0].rows:
        k1 = row.cells[0].text.strip()
        if k1 in draft.meta:
            _set_cell(row.cells[1], draft.meta[k1])
        if len(row.cells) > 2:
            k2 = row.cells[2].text.strip()
            if k2 in draft.meta:
                _set_cell(row.cells[3], draft.meta[k2])

    if len(doc.paragraphs) > 1:
        _set_para(doc.paragraphs[1], draft.position_title)

    if doc.tables and len(doc.tables) > 1 and doc.tables[1].rows:
        _set_cell(
            doc.tables[1].rows[0].cells[0],
            f"Назначение должности. {draft.goal_summary}",
        )

    idx = 0
    for pi, para in enumerate(doc.paragraphs):
        if pi == 1:
            continue
        text = para.text.strip()
        if not text or text.startswith(SKIP_HEADINGS):
            continue
        if idx < len(draft.content_paragraphs):
            _set_para(para, draft.content_paragraphs[idx])
            idx += 1

    if len(doc.tables) > 2 and draft.cycle_rows:
        tbl = _find_table_by_header(doc, "период", "основные действия") or doc.tables[2]
        _fill_table(tbl, draft.cycle_rows)
    if len(doc.tables) > 3 and draft.kpi_rows:
        tbl = _find_table_by_header(doc, "показатель", "как измеряется") or doc.tables[3]
        _fill_table(tbl, draft.kpi_rows)
    approvers_tbl = _find_table_by_header(doc, "роль", "фio", "подпись")
    if approvers_tbl:
        _fill_table(
            approvers_tbl,
            [
                ["Роль", "ФИО / подпись", "Примечание"],
                ["Руководитель", "__________________", "Согласует"],
                ["HR / методолог", "__________________", "Проверяет структуру"],
            ],
        )
    elif len(doc.tables) > 4:
        _fill_table(
            doc.tables[4],
            [
                ["Роль", "ФИО / подпись", "Примечание"],
                ["Руководитель", "__________________", "Согласует"],
                ["HR / методолог", "__________________", "Проверяет структуру"],
            ],
        )
    skills_tbl = _find_table_by_header(doc, "навык")
    if skills_tbl and draft.skill_rows and len(draft.skill_rows) > 1:
        _fill_table(skills_tbl, draft.skill_rows)

    doc.save(str(output_path))
    return output_path


def generate_regulation_from_web(
    position_title: str,
    comment: str | None = None,
    *,
    template_code: str = "default",
) -> tuple[Path, RegulationDraft]:
    del template_code  # reserved for future template-specific layouts
    query = f"должностная инструкция регламент {position_title} обязанности KPI"
    if comment and comment.strip():
        query = f"{query} {comment.strip()[:120]}"
    snippets = search_web_snippets(query)
    draft = compose_draft(position_title, comment, snippets)

    safe_name = re.sub(r"[^\w\-]+", "_", draft.regulation_code, flags=re.UNICODE)
    filename = f"Регламент_{safe_name}_{uuid.uuid4().hex[:8]}.docx"
    output_path = GENERATED_DIR / filename
    render_docx(draft, output_path)
    return output_path, draft
